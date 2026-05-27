import streamlit as st
import os
from io import BytesIO
from docx import Document
from langchain_openai import ChatOpenAI
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.prompts import ChatPromptTemplate

st.set_page_config(page_title="Generador de Evaluaciones", page_icon="📝", layout="wide")

st.title("📝 Generador Automatizado de Exámenes y Rúbricas")
st.write("Sube una lectura o capítulo de libro para estructurar un banco de preguntas técnico y su rúbrica analítica.")

# 1. Validación de Credenciales en Render
if "OPENAI_API_KEY" in os.environ:
    api_key_render = os.getenv("OPENAI_API_KEY")
else:
    st.error("Por favor, configura la variable OPENAI_API_KEY en el panel de Render.")
    st.stop()

# 2. Configuración de Parámetros Pedagógicos en la barra lateral
with st.sidebar:
    st.header("⚙️ Parámetros del Instrumento")
    
    tipo_pregunta = st.selectbox(
        "Tipo de preguntas:",
        ["Opción Múltiple (con justificación)", "Análisis de Caso / Ensayo", "Preguntas de Desarrollo Técnico"]
    )
    
    nivel_bloom = st.selectbox(
        "Nivel cognitivo (Taxonomía de Bloom):",
        ["Recordar y Comprender (Básico)", "Aplicar y Analizar (Intermedio)", "Evaluar y Crear (Avanzado)"]
    )
    
    cantidad_preguntas = st.slider("Cantidad de preguntas a generar:", min_value=3, max_value=10, value=5)
    
    criterios_rubrica = st.text_input(
        "Criterios de la rúbrica (separados por comas):",
        value="Rigor conceptual, Capacidad de análisis, Coherencia argumentativa"
    )

# 3. Procesamiento del PDF de origen
archivo_pdf = st.file_uploader("Sube el material base (PDF con texto seleccionable)", type="pdf")

def transformar_a_word(texto_examen):
    """Función para volcar la respuesta en un documento Word real"""
    doc = Document()
    doc.add_heading("BANCO DE PREGUNTAS E INSTRUMENTO DE EVALUACIÓN", level=1)
    
    for linea in texto_examen.split("\n"):
        if linea.startswith("###"):
            doc.add_heading(linea.replace("###", "").strip(), level=3)
        elif linea.startswith("##"):
            doc.add_heading(linea.replace("##", "").strip(), level=2)
        elif linea.startswith("#"):
            doc.add_heading(linea.replace("#", "").strip(), level=1)
        else:
            doc.add_paragraph(linea)
            
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

if archivo_pdf:
    with open("temp_eval.pdf", "wb") as f:
        f.write(archivo_pdf.getbuffer())
        
    st.info("📖 Extrayendo y analizando el contenido del documento...")
    loader = PyPDFLoader("temp_eval.pdf")
    docs = loader.load()
    docs_con_texto = [d for d in docs if d.page_content.strip()]
    
    if not docs_con_texto:
        st.error("No se pudo extraer texto del PDF. Asegúrate de que no sea un escaneo.")
        st.stop()
        
    # Unificamos las primeras páginas para el contexto
    texto_contexto = "\n".join([d.page_content for d in docs_con_texto[:6]])

    if st.button("🚀 Generar Evaluación y Rúbrica"):
        with st.spinner("La IA está analizando los conceptos y redactando el instrumento bajo el enfoque por competencias..."):
            
            # Usamos gpt-4o-mini (más inteligente y económico)
            llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.3, openai_api_key=api_key_render)
            
            template_evaluacion = (
                "Eres un especialista en diseño curricular y evaluación bajo el enfoque de competencias universitarias.\n"
                "Tu tarea es analizar el siguiente texto de origen y diseñar un banco de preguntas junto con una rúbrica analítica.\n\n"
                "--- TEXTO DE ORIGEN ---\n"
                "{contexto}\n"
                "------------------------\n\n"
                "INSTRUCCIONES DE DISEÑO:\n"
                "1. Genera exactamente {cantidad} preguntas de tipo: '{tipo}'.\n"
                "2. Alinea el nivel de las preguntas al nivel cognitivo de Bloom: '{nivel}'.\n"
                "3. Las preguntas deben exigir el uso explícito del texto adjunto, evitando la memorización superficial.\n"
                "4. Si elegiste Opción Múltiple, incluye las alternativas (A, B, C, D) e indica la respuesta correcta justificando pedagógicamente por qué las demás son distractores.\n"
                "5. Diseña una MATRIZ DE RÚBRICA ANALÍTICA al final del documento considerando exclusivamente estos criterios de evaluación: {criterios}.\n"
                "6. Para cada criterio, detalla los niveles de desempeño: Sobresaliente (4 pts), Logrado (3 pts), En Proceso (2 pts) e Inicio (1 pt).\n\n"
                "FORMATO DE SALIDA:\n"
                "Usa títulos claros utilizando sintaxis Markdown con '#' y '##' para separar las secciones de manera muy profesional. No inventes datos fuera de la lectura."
            )
            
            prompt = ChatPromptTemplate.from_template(template_evaluacion)
            cadena_procesamiento = prompt | llm
            
            respuesta_ia = cadena_procesamiento.invoke({
                "contexto": texto_contexto,
                "cantidad": cantidad_preguntas,
                "tipo": tipo_pregunta,
                "nivel": nivel_bloom,
                "criterios": criterios_rubrica
            })
            
            resultado_texto = respuesta_ia.content
            
            st.success("¡Instrumento generado con éxito!")
            st.markdown(resultado_texto)
            
            archivo_word = transformar_a_word(resultado_texto)
            
            st.download_button(
                label="📥 Descargar Examen y Rúbrica (.docx)",
                data=archivo_word,
                file_name="Evaluacion_y_Rubrica_Generada.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
